import scrapy
from scrapy.spiders import CrawlSpider, Rule
from scrapy.linkextractors import LinkExtractor
import docx
from copy import deepcopy
from docx.shared import Pt
from scholarly import scholarly
from sqlalchemy.orm import sessionmaker

from ..database import db
import re


Session = sessionmaker(bind = db.engine)
session = Session()

def removeAccent(s):
    s = re.sub('[áàảãạăắằẳẵặâấầẩẫậ]', 'a', s)
    s = re.sub('[ÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬ]', 'A', s)
    s = re.sub('[éèẻẽẹêếềểễệ]', 'e', s)
    s = re.sub('[ÉÈẺẼẸÊẾỀỂỄỆ]', 'E', s)
    s = re.sub('[óòỏõọôốồổỗộơớờởỡợ]', 'o', s)
    s = re.sub('[ÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢ]', 'O', s)
    s = re.sub('[íìỉĩị]', 'i', s)
    s = re.sub('[ÍÌỈĨỊ]', 'I', s)
    s = re.sub('[úùủũụưứừửữự]', 'u', s)
    s = re.sub('[ÚÙỦŨỤƯỨỪỬỮỰ]', 'U', s)
    s = re.sub('[ýỳỷỹỵ]', 'y', s)
    s = re.sub('[ÝỲỶỸỴ]', 'Y', s)
    s = re.sub('đ', 'd', s)
    s = re.sub('Đ', 'D', s)
    return s

def addText(container, string, style, clear = False):
    if clear:
        container.text = ""
    paragraph = container.add_paragraph('', style = style)
    paragraph.add_run(string).italic = True
    return 


class Teacher(scrapy.Item):
    name = scrapy.Field()
    soictLink = scrapy.Field()
    ggscholarLink = scrapy.Field()

    designation = scrapy.Field()
    academicCareer = scrapy.Field()

    publicationsOverLast5Years = scrapy.Field()   

    proprietary = scrapy.Field()

    link = scrapy.Field()

class AIOSpider(CrawlSpider):
    name = "test"
    base_url = 'soict.hust.edu.vn'
    le = LinkExtractor(restrict_xpaths =[ './/h2[@class = "entry-title no-margin"]/a'] ) 
    start_urls = [
            'https://soict.hust.edu.vn/pgs-la-the-vinh.html'
        ]
    def start_requests(self):
        for url in self.start_urls:
            yield scrapy.Request(url=url, callback=self.parse_soict)

    def parse_soict(self, response):

        #-----------------------------------------#Save item values
        item = Teacher()

        item['name'] = response.xpath('//p[@class = "lead"]/span/strong/text()').get()
        if item['name'] == None : 
            item['name'] = response.xpath('//p[@class = "lead"]/strong/span/text()').get()
        if item['name'] == None : 
            item['name'] = response.xpath('//h1/strong/text()').get()
        if item['name'] == None : 
            item['name'] = response.xpath('//p[@class = "lead"]/span/b/text()').get()

        try:
            item['designation'] = response.xpath('//p/strong/text()')
            if item['designation'] == [] :
                item['designation'] = response.xpath('//p/b/text()')
            if item['designation'] == [] :
                item['designation'] = response.xpath('//h4/strong/text()')
            if item['designation'] == [] :
                item['designation'] = response.xpath('//p/strong/span/text()')
        except:
            pass

        try:
            if item['name'] not in ['Nguyễn An Hưng', 'Michel Toulouse']:
                item['academicCareer'] = []
                temp = response.xpath('//div[@class = "col-inner"]')[0]
                i = 0
                while True:
                    s = temp.xpath("//p/text()").getall()[i].replace('\n','')
                    if ( ("Email" not in s) and 'muriel' not in s) :
                        if s != "":
                            item['academicCareer'].append(s)
                        i+=1
                    else:
                        break
            else:
                temp = response.xpath('//div[@class = "col-inner"]/p[@class = "lead"]/following-sibling::div/text()')
                item['academicCareer'] = [s for s in temp.getall() if "Email" not in s]
        except:
            pass

        try:
            spans = response.xpath('//div[@class = "container section-title-container"]//h3//span[text() = "Bằng sáng chế"]/ancestor::div[@class = "container section-title-container"]/following-sibling::ul[1]/li/text()')
            item['proprietary'] = [span.get() for span in spans]
        except:
            pass
        item['link'] = response.request.url


        yield item
        #-----------------------------------------#


        #-----------------------------------------#Init output files
        try:
            outputFile = docx.Document(f'testoutput/{item["name"]}.docx')
        except:
            inputFile = docx.Document('Staff handbook template.docx')
            copyContent = deepcopy(inputFile)
            copyContent.save(f'testoutput/{item["name"]}.docx')
            outputFile = docx.Document(f'testoutput/{item["name"]}.docx')


        style = outputFile.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(10)

        #-----------------------------------------#


        #-----------------------------------------# Name and designation
        addText(outputFile.tables[0].columns[1].cells[0], item['name'], style = style, clear = True )
        outputFile.tables[0].columns[1].cells[1].text = ''
        for d in item["designation"]:
            addText(outputFile.tables[0].columns[1].cells[1], d.get(), style = style)

        
        #-----------------------------------------#


        #-----------------------------------------# Academic Career
        outputFile.tables[0].columns[1].cells[2].text = ''
        outputFile.tables[0].columns[5].cells[2].text = ''
        outputFile.tables[0].columns[7].cells[2].text = ''

        for s in item['academicCareer']:
            first_bracket_index = s.find("(")

            subject_institution_year = s[first_bracket_index+1:s.find(")")].split(', ')
            if(len(subject_institution_year) >= 3):
                addText(outputFile.tables[0].columns[1].cells[2],f"{s[:first_bracket_index]}, {subject_institution_year[0]}",
                      style = style)
                addText(outputFile.tables[0].columns[5].cells[2], subject_institution_year[1], style = style)
            else:
                addText(outputFile.tables[0].columns[1].cells[2],f"{s[:first_bracket_index]}",
                      style = style)
                addText(outputFile.tables[0].columns[5].cells[2], subject_institution_year[0], style = style)
            addText(outputFile.tables[0].columns[7].cells[2], subject_institution_year[-1], style = style)
        #-----------------------------------------#


        #-----------------------------------------# Proprietary
        addText(outputFile.tables[0].columns[1].cells[6], "Title" , style = style, clear = True)
        addText(outputFile.tables[0].columns[6].cells[6], "Year" , style = style, clear = True)

        for s in item['proprietary']:
            first_quotation_index = s.find('“')
                
            title = s[first_quotation_index + 1 : s.find('”', first_quotation_index + 1)]
            year = s[-5:-1]

            self.prop_id += 1
            prop = db.Proprietary(self.prop_id, title, year, self.teacher_id)
            self.session.add(prop)

            addText(outputFile.tables[0].columns[1].cells[6], title, style = style)
            addText(outputFile.tables[0].columns[6].cells[6], year, style = style)
        #-----------------------------------------#


        #-----------------------------------------#Save output files
        outputFile.save(f'testoutput/{item["name"]}.docx')
        #-----------------------------------------#

        search_query = scholarly.search_author(removeAccent(item['name']))

        outputFile.tables[0].columns[1].cells[7].text = ''
            
        while True:
            author = scholarly.fill(next(search_query))
            if author['email_domain'] == "@soict.hust.edu.vn":
                print(item['name'] + ' found')
                #writePubs(outputFile, author['publications'], style)
                publications = author['publications']
                numPubs = len(publications)
                try:
                    numRecentPubs = f"Selected recent publications from a total of approx. : {numPubs}"
                    outputFile.tables[0].columns[1].cells[7].text = ""
                except:
                    return
                for pub in publications:
                    try:
                        if int(pub['bib']['pub_year']) >= 2018:
                            pub = scholarly.fill(pub)

                            try:
                                author = f"{pub['bib']['author']}"
                            except:
                                author = ""
                            try:
                                title = f"{pub['bib']['title']}"
                            except:
                                title =  ""
                            try:
                                otherInfo = f"{pub['pub_url']}"
                            except:
                                otherInfo = f""
                            try:
                                publisher = f"{pub['bib']['publisher']}"
                            except:
                                publisher = f""
                            try:
                                dateOfPublication = f"{pub['bib']['pub_year']}"
                            except: 
                                dateOfPublication = f""
                            try:
                                publicationJournal = f"Journal {pub['bib']['journal']}, volume {pub['bib']['volume']}, page {pub['bib']['pages']}"
                            except:
                                publicationJournal = f""
                            addText(outputFile.tables[0].columns[1].cells[7] , numRecentPubs , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Author(s) : " + author , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Title : " + title , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Any other information : " + otherInfo , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Publisher : " + publisher , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Year of publication : " + dateOfPublication , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , publicationJournal , style )
                            
                    except:
                        continue
                break
        print(item["name"] + " saved")



