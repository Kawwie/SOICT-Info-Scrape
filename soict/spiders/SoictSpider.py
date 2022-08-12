import scrapy
from scrapy.spiders import CrawlSpider, Rule
from scrapy.linkextractors import LinkExtractor
import docx
from copy import deepcopy

from docx.shared import Pt

from ..database import db
from sqlalchemy.orm import sessionmaker

import re



def addText(container, string, style, clear = False):
    if clear:
        container.text = ""
    paragraph = container.add_paragraph('', style = style)
    paragraph.add_run(string).italic = True
    return 

class Teacher(scrapy.Item):
    name = scrapy.Field()
    soictLink = scrapy.Field()
    ggscholarLink = scrapy.Field()

    designation = scrapy.Field()
    academicCareer = scrapy.Field()

    publicationsOverLast5Years = scrapy.Field()   

    proprietary = scrapy.Field()

    link = scrapy.Field()

class SoictSpider(CrawlSpider):


    name = "soict"
    base_url = 'soict.hust.edu.vn'
    start_urls = [
            'https://soict.hust.edu.vn/can-bo/',
            'https://soict.hust.edu.vn/can-bo/page/2',
            'https://soict.hust.edu.vn/can-bo/page/3',
            'https://soict.hust.edu.vn/can-bo/page/4',
            'https://soict.hust.edu.vn/can-bo/page/5',
        ]

    rules = [
    Rule(LinkExtractor(restrict_xpaths =[ './/h2[@class = "entry-title no-margin"]/a'] , 
                       allow_domains = ["soict.hust.edu.vn"] ), 
                        callback='parse_soict', 
                        follow = False )

    ]
    
    #-----------------------------------------#Initialize variables
    Session = sessionmaker(db.engine)
    session = Session()
    teacher_id = 0
    prop_id = 0
    #-----------------------------------------#


    def parse_soict(self, response):

        #-----------------------------------------#Save item values
        item = Teacher()

        item['name'] = response.xpath('//p[@class = "lead"]/span/strong/text()').get()
        if item['name'] == None : 
            item['name'] = response.xpath('//p[@class = "lead"]/strong/span/text()').get()
        if item['name'] == None : 
            item['name'] = response.xpath('//h1/strong/text()').get()
        if item['name'] == None : 
            item['name'] = response.xpath('//p[@class = "lead"]/span/b/text()').get()


        item['designation'] = response.xpath('//p/strong/text()')
        if item['designation'] == [] :
            item['designation'] = response.xpath('//p/b/text()')
        if item['designation'] == [] :
            item['designation'] = response.xpath('//h4/strong/text()')
        if item['designation'] == [] :
            item['designation'] = response.xpath('//p/strong/span/text()')
        
        if item['name'] not in ['Nguyễn An Hưng', 'Michel Toulouse']:
            item['academicCareer'] = []
            temp = response.xpath('//div[@class = "col-inner"]')[0]
            i = 0
            while True:
                s = temp.xpath("//p/text()").getall()[i].replace('\n','')
                if ( ("Email" not in s) ) :
                    if s != "":
                        item['academicCareer'].append(s)
                    i+=1
                else:
                    break
        else:
            temp = response.xpath('//div[@class = "col-inner"]/p[@class = "lead"]/following-sibling::div/text()')
            item['academicCareer'] = [s.get() for s in temp if "Email" not in s]

        try:
            spans = response.xpath('//div[@class = "container section-title-container"]//h3//span[text() = "Bằng sáng chế"]/ancestor::div[@class = "container section-title-container"]/following-sibling::ul[1]/li/text()')
            item['proprietary'] = [span.get() for span in spans]
        except:
            pass
        item['link'] = response.request.url
        #yield item
        #-----------------------------------------#


        #-----------------------------------------#Init output files
        try:
            outputFile = docx.Document(f'output/{item["name"]}.docx')
        except:
            inputFile = docx.Document('Staff handbook template.docx')
            copyContent = deepcopy(inputFile)
            copyContent.save(f'output/{item["name"]}.docx')
            outputFile = docx.Document(f'output/{item["name"]}.docx')


        style = outputFile.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(10)

        #-----------------------------------------#


        #-----------------------------------------# Name and designation
        addText(outputFile.tables[0].columns[1].cells[0], item['name'], style = style, clear = True )
        outputFile.tables[0].columns[1].cells[1].text = ''
        for d in item["designation"]:
            addText(outputFile.tables[0].columns[1].cells[1], d.get(), style = style)

        
        #-----------------------------------------#


        #-----------------------------------------# Academic Career
        outputFile.tables[0].columns[1].cells[2].text = ''
        outputFile.tables[0].columns[5].cells[2].text = ''
        outputFile.tables[0].columns[7].cells[2].text = ''

        for s in item['academicCareer']:
            first_bracket_index = s.find("(")

            subject_institution_year = s[first_bracket_index+1:s.find(")")].split(', ')
            if(len(subject_institution_year) >= 3):
                addText(outputFile.tables[0].columns[1].cells[2],f"{s[:first_bracket_index]}, {subject_institution_year[0]}",
                      style = style)
                addText(outputFile.tables[0].columns[5].cells[2], subject_institution_year[1], style = style)
            else:
                addText(outputFile.tables[0].columns[1].cells[2],f"{s[:first_bracket_index]}",
                      style = style)
                addText(outputFile.tables[0].columns[5].cells[2], subject_institution_year[0], style = style)
            addText(outputFile.tables[0].columns[7].cells[2], subject_institution_year[-1], style = style)
        #-----------------------------------------#

        self.teacher_id += 1
        teacher = db.Teacher(self.teacher_id, item['name'], '\n'.join([d.get() for d in item['designation']]), '\n'.join([c for c in item['academicCareer']]), item['link'])
        self.session.add(teacher)


        #-----------------------------------------# Proprietary
        addText(outputFile.tables[0].columns[1].cells[6], "Title" , style = style, clear = True)
        addText(outputFile.tables[0].columns[6].cells[6], "Year" , style = style, clear = True)

        for s in item['proprietary']:
            first_quotation_index = s.find('“')
                
            title = s[first_quotation_index + 1 : s.find('”', first_quotation_index + 1)]
            year = s[-5:-1]

            self.prop_id += 1
            prop = db.Proprietary(self.prop_id, title, year, self.teacher_id)
            self.session.add(prop)

            addText(outputFile.tables[0].columns[1].cells[6], title, style = style)
            addText(outputFile.tables[0].columns[6].cells[6], year, style = style)
        #-----------------------------------------#


        #-----------------------------------------#Save output files
        outputFile.save(f'output/{item["name"]}.docx')
        #-----------------------------------------#



        #-----------------------------------------#Commit to DB
        self.session.commit()
        #-----------------------------------------#

        print(item["name"] + " saved")
