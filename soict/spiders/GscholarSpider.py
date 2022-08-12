import scrapy
from scrapy.spiders import CrawlSpider, Rule
from scrapy.linkextractors import LinkExtractor
import docx
from copy import deepcopy

from scholarly import scholarly
from docx.shared import Pt

from ..database import db
from sqlalchemy.orm import sessionmaker

import re

def removeAccent(s):
    s = re.sub('[áàảãạăắằẳẵặâấầẩẫậ]', 'a', s)
    s = re.sub('[ÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬ]', 'A', s)
    s = re.sub('[éèẻẽẹêếềểễệ]', 'e', s)
    s = re.sub('[ÉÈẺẼẸÊẾỀỂỄỆ]', 'E', s)
    s = re.sub('[óòỏõọôốồổỗộơớờởỡợ]', 'o', s)
    s = re.sub('[ÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢ]', 'O', s)
    s = re.sub('[íìỉĩị]', 'i', s)
    s = re.sub('[ÍÌỈĨỊ]', 'I', s)
    s = re.sub('[úùủũụưứừửữự]', 'u', s)
    s = re.sub('[ÚÙỦŨỤƯỨỪỬỮỰ]', 'U', s)
    s = re.sub('[ýỳỷỹỵ]', 'y', s)
    s = re.sub('[ÝỲỶỸỴ]', 'Y', s)
    s = re.sub('đ', 'd', s)
    s = re.sub('Đ', 'D', s)
    return s

def addText(container, string, style, clear = False):
    if clear:
        container.text = ""
    paragraph = container.add_paragraph('', style = style)
    paragraph.add_run(string).italic = True
    return 



class GscholarSpider(CrawlSpider):
    name = "GS"
    base_url = 'soict.hust.edu.vn'
    le = LinkExtractor(restrict_xpaths =[ './/h2[@class = "entry-title no-margin"]/a'] ) #restrict_xpaths =[ './/h2[@class = "entry-title no-margin"]/a']
    start_urls = [
            'https://soict.hust.edu.vn/can-bo/',
            'https://soict.hust.edu.vn/can-bo/page/2',
            'https://soict.hust.edu.vn/can-bo/page/3',
            'https://soict.hust.edu.vn/can-bo/page/4',
            'https://soict.hust.edu.vn/can-bo/page/5',
        ]

    rules = [
    Rule(LinkExtractor(restrict_xpaths =[ './/h2[@class = "entry-title no-margin"]/a'] , 
                       allow_domains = ["soict.hust.edu.vn"] ), 
                        callback='parse_soict', 
                        follow = False )

    ]
    
    #-----------------------------------------#Initialize variables
    Session = sessionmaker(db.engine)
    session = Session()
    pub_id = 0
    #-----------------------------------------#

    def parse_soict(self, response):

        name = response.xpath('//p[@class = "lead"]/span/strong/text()').get()  


        try:
            outputFile = docx.Document(f'output/{name}.docx')
        except:
            inputFile = docx.Document('Staff handbook template.docx')
            copyContent = deepcopy(inputFile)
            copyContent.save(f'output/{name}.docx')
            outputFile = docx.Document(f'output/{name}.docx')

        style = outputFile.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(10)


        print(name)

        search_query = scholarly.search_author(removeAccent(name))

        outputFile.tables[0].columns[1].cells[7].text = ''
            
        while True:
            author = scholarly.fill(next(search_query))
            if author['email_domain'] == "@soict.hust.edu.vn":
                print(name + ' found')
                #writePubs(outputFile, author['publications'], style)
                publications = author['publications']
                numPubs = len(publications)
                try:
                    numRecentPubs = f"Selected recent publications from a total of approx. : {numPubs}"
                    outputFile.tables[0].columns[1].cells[7].text = ""
                except:
                    return
                for pub in publications:
                    try:
                        if int(pub['bib']['pub_year']) >= 2018:
                            pub = scholarly.fill(pub)

                            try:
                                author = f"{pub['bib']['author']}"
                            except:
                                author = ""
                            try:
                                title = f"{pub['bib']['title']}"
                            except:
                                title =  ""
                            try:
                                otherInfo = f"{pub['pub_url']}"
                            except:
                                otherInfo = f""
                            try:
                                publisher = f"{pub['bib']['publisher']}"
                            except:
                                publisher = f""
                            try:
                                dateOfPublication = f"{pub['bib']['pub_year']}"
                            except: 
                                dateOfPublication = f""
                            try:
                                publicationJournal = f"Journal {pub['bib']['journal']}, volume {pub['bib']['volume']}, page {pub['bib']['pages']}"
                            except:
                                publicationJournal = f""
                            addText(outputFile.tables[0].columns[1].cells[7] , numRecentPubs , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Author(s) : " + author , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Title : " + title , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Any other information : " + otherInfo , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Publisher : " + publisher , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , "Year of publication : " + dateOfPublication , style )
                            addText(outputFile.tables[0].columns[1].cells[7] , publicationJournal , style )
                            
                            self.pub_id += 1
                            teacher_id = self.session.query(db.Teacher).filter(db.Teacher.name == name)[0].teacher_id
                            pub = db.Publication(self.pub_id, title, author, otherInfo, publisher, dateOfPublication, publicationJournal, teacher_id)
                            self.session.add(pub)
                            self.session.commit()
                    except:
                        continue
                break

        outputFile.save(f'output/{name}.docx')
        print(name + " saved")
